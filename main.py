from docx import Document
import re
from secondPage import remove_text_in_brackets


exportFileName = "text"
def expand_ranges(text):
    """
    Преобразует числовые диапазоны вида "01 - 03", "101 - 103", "1001 - 1003" в список чисел.
    """
    def replace_range(match):
        start, end = map(int, match.group(1, 2))
        width = len(match.group(1))  # Длина числа (2, 3, 4 знака)
        return ", ".join(f"{i:0{width}}" for i in range(start, end + 1))

    # Изменённое регулярное выражение для чисел с длиной от 2 до 4 символов
    return re.sub(r'(\d{2,4})\s*-\s*(\d{2,4})', replace_range, text)


def process_docx(file_path, exportFileName):
    """
    Анализирует и редактирует docx файл в соответствии с требованиями.
    """
    # Открытие документа
    file_path = "vpo1.docx"

    doc = Document(file_path)

    # Новый список для хранения обработанных абзацев
    processed_paragraphs = []

    for paragraph in doc.paragraphs:
        text = paragraph.text
        is_bold = any(run.bold for run in paragraph.runs)

        # Если текст содержит "по строкам" или "По строкам"
        if "по графам" in text.lower():
            processed_paragraphs.append(expand_ranges(text))
        # Если текст содержит "по строке" или "По строке"
        elif "по графе" in text.lower():
            processed_paragraphs.append(text)
        # Если текст не содержит ключевые слова, но есть жирный шрифт, сохраняем
        elif "в графе" in text.lower():
            processed_paragraphs.append(text)
        elif "в графах" in text.lower():
            processed_paragraphs.append(expand_ranges(text))
        elif is_bold:
            processed_paragraphs.append(text)
        # В противном случае игнорируем абзац

    # Очистка текущего содержимого документа
    for paragraph in doc.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    # Добавление обработанных абзацев в документ
    for p_text in processed_paragraphs:
        doc.add_paragraph(p_text)

    # Сохранение документа
    exportFileName = ("processed_" + file_path)
    ds = doc.save("processed_" + file_path)
    return exportFileName



# Пример использования
exportFileName = process_docx("gg", exportFileName)
print(exportFileName)
remove_text_in_brackets(exportFileName, "processed_inov.docx")
